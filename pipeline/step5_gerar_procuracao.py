import os
import tempfile
import subprocess
import shutil
import platform
from datetime import datetime, timezone, timedelta
from pathlib import Path
import docx

_MESES_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
}

# Preposições/conjunções que ficam em minúsculo em nomes próprios em PT-BR
_PALAVRAS_MINUSCULAS = {"de", "da", "do", "das", "dos", "e"}

# Fuso horário de Brasília (BRT = UTC-3)
_TZ_BR = timezone(timedelta(hours=-3))


def _data_por_extenso_pt() -> str:
    """Retorna a data atual em português no fuso de Brasília: '17 de abril de 2026'."""
    hoje = datetime.now(_TZ_BR)
    return f"{hoje.day} de {_MESES_PT[hoje.month]} de {hoje.year}"


def _titulo_pt(texto: str) -> str:
    """Capitaliza nome próprio em PT-BR mantendo preposições em minúsculo.

    Ex: 'RIO DE JANEIRO' → 'Rio de Janeiro'
        'SÃO JOSÉ DOS CAMPOS' → 'São José dos Campos'
    A primeira palavra fica sempre capitalizada.
    """
    partes = texto.strip().split()
    if not partes:
        return ""
    resultado = [partes[0].capitalize()]
    for p in partes[1:]:
        if p.lower() in _PALAVRAS_MINUSCULAS:
            resultado.append(p.lower())
        else:
            resultado.append(p.capitalize())
    return " ".join(resultado)

def _get_soffice_cmd() -> str:
    """Retorna o executável do LibreOffice"""
    if platform.system() == "Windows":
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        for c in candidates:
            if os.path.exists(c):
                return c
        return "soffice.exe"
    return "soffice"

def gerar_procuracao_pdf(dados, pasta_saida: str) -> str:
    docx_base = str(Path(__file__).parent.parent / "PROCURAÇÃO ENERGISA MT - CERÂMICA PROGRESSO.docx")
    if not os.path.exists(docx_base):
        raise FileNotFoundError(f"Arquivo base não encontrado: {docx_base}")
    
    # 1. Copiar pro temp
    tmp_dir = tempfile.mkdtemp(prefix="step5_procuracao_")
    tmp_docx = os.path.join(tmp_dir, "procuracao_temp.docx")
    shutil.copy2(docx_base, tmp_docx)

    # 2. Modificar docx
    doc = docx.Document(tmp_docx)
    
    # Montar Endereço
    endereco = f"{dados.logradouro}"
    if dados.numero:
        endereco += f", {dados.numero}"
    if dados.bairro:
        endereco += f", {dados.bairro}"
    if dados.cidade and dados.uf:
        endereco += f", {dados.cidade}-{dados.uf}"
    endereco = endereco.upper()

    # Substituições do OUTORGANTE (titular/cliente)
    subs = {
        "MARCOS ANTONIO GOMES": (dados.titular or "").upper(),
        "298.607.681-53": dados.cpf_cnpj or "",
        "R.SICILIA, SN, RESICENCIAL FLORENÇA, SINOP-MT": endereco,
    }

    # Data de emissão (substitui a linha "Sinop, 31 de março de 2026")
    cidade_proc = _titulo_pt(dados.cidade or "Sinop")
    data_hoje = _data_por_extenso_pt()
    subs["Sinop, 31 de março de 2026"] = f"{cidade_proc}, {data_hoje}"

    # Substituições do OUTORGADO (engenheiro/responsável técnico)
    resp_nome = getattr(dados, "resp_nome", "") or ""
    resp_cpf = getattr(dados, "resp_cpf", "") or ""
    resp_endereco = getattr(dados, "resp_endereco", "") or ""
    if resp_nome.strip():
        subs["Edimo Perondi Junior"] = resp_nome.strip()
    if resp_cpf.strip():
        subs["058.029.991-01"] = resp_cpf.strip()
    if resp_endereco.strip():
        subs["Rua Giuliana, 1105, apartamento 04, Residencial Florença, Sinop, Mato Grosso"] = resp_endereco.strip()

    def _replace_in_paragraph(para, old: str, new: str):
        if old not in para.text:
            return
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return
        # Fallback: texto dividido em múltiplos runs — reconstrói preservando o 1º run
        if para.runs:
            para.runs[0].text = para.text.replace(old, new)
            for run in para.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        for old, new in subs.items():
            _replace_in_paragraph(p, old, new)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in subs.items():
                        _replace_in_paragraph(p, old, new)

    doc.save(tmp_docx)
    
    # 3. Converter para PDF
    nome_pdf = f"{(dados.titular or '').upper().strip()}_UC_{dados.codigo_uc}_PROCURACAO.pdf"
    caminho_pdf = os.path.join(pasta_saida, nome_pdf)
    Path(pasta_saida).mkdir(parents=True, exist_ok=True)
    
    soffice = _get_soffice_cmd()
    cmd = [
        soffice,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", pasta_saida,
        tmp_docx
    ]
    
    print(f"  [step5] Convertendo docx para pdf...")
    subprocess.run(cmd, capture_output=True, check=True)
    
    # Libreoffice salva como procuracao_temp.pdf
    lo_pdf = os.path.join(pasta_saida, "procuracao_temp.pdf")
    if os.path.exists(lo_pdf):
        if os.path.exists(caminho_pdf):
            os.remove(caminho_pdf)
        os.rename(lo_pdf, caminho_pdf)
    
    shutil.rmtree(tmp_dir, ignore_errors=True)
    
    if os.path.exists(caminho_pdf):
        print(f"  [step5] OK — Procuração em PDF gerada: {nome_pdf}")
        return caminho_pdf
    else:
        raise RuntimeError("Falha ao gerar o PDF da procuração (LibreOffice não criou o arquivo).")
